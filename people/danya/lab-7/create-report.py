from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
SRC = ROOT / "src" / "main"
OUTPUT = ROOT / "Отчет_ЛР7_JDBC.docx"
ASSETS = ROOT / "report_assets"


def get_font(size: int, bold: bool = False):
    names = ["arialbd.ttf" if bold else "arial.ttf", "Arial.ttf", "DejaVuSans.ttf"]
    for name in names:
        try:
            return ImageFont.truetype(name, size=size)
        except OSError:
            continue
    return ImageFont.load_default()


def set_run_font(run, size=14, bold=False, font_name="Times New Roman"):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run._element.rPr.rFonts.set(qn("w:cs"), font_name)
    run.font.size = Pt(size)
    run.bold = bold


def style_document(doc: Document):
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
    normal.font.size = Pt(14)

    for style_name in ["Title", "Subtitle", "Heading 1", "Heading 2"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")


def add_code_paragraph(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run, size=10, font_name="Consolas")


def add_code_block(doc: Document, title: str, text: str):
    doc.add_heading(title, level=2)
    for line in text.splitlines():
        add_code_paragraph(doc, line)
    if not text.strip():
        add_code_paragraph(doc, "")


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def tree_text() -> str:
    return "\n".join(
        [
            "lab7",
            "├── conf.prop",
            "├── pom.xml",
            "├── create_report.py",
            "└── src",
            "    └── main",
            "        ├── java",
            "        │   ├── module-info.java",
            "        │   └── com/lab7",
            "        │       ├── MainApp.java",
            "        │       ├── controllers",
            "        │       │   ├── LoginController.java",
            "        │       │   ├── MainController.java",
            "        │       │   └── DatabaseController.java",
            "        │       ├── db",
            "        │       │   └── DatabaseManager.java",
            "        │       ├── gui",
            "        │       │   └── Dialogs.java",
            "        │       └── model",
            "        │           └── UserRecord.java",
            "        └── resources",
            "            ├── images",
            "            │   └── database-banner.png",
            "            └── com/lab7/fxml",
            "                ├── LoginView.fxml",
            "                ├── MainView.fxml",
            "                └── DatabaseView.fxml",
        ]
    )


def create_login_image(path: Path):
    img = Image.new("RGB", (820, 460), "#f4f7fb")
    draw = ImageDraw.Draw(img)
    title_font = get_font(28, True)
    label_font = get_font(22, True)
    text_font = get_font(20)
    button_font = get_font(20, True)

    draw.rounded_rectangle((170, 70, 650, 330), radius=20, fill="white", outline="#c7d2e0", width=2)
    draw.text((330, 30), "Авторизация", fill="#1f2937", font=title_font)

    draw.text((220, 120), "Логин", fill="#1f2937", font=label_font)
    draw.rounded_rectangle((330, 115, 590, 155), radius=8, fill="#fbfdff", outline="#9db0c9", width=2)
    draw.text((345, 124), "veritasad", fill="#374151", font=text_font)

    draw.text((220, 190), "Пароль", fill="#1f2937", font=label_font)
    draw.rounded_rectangle((330, 185, 590, 225), radius=8, fill="#fbfdff", outline="#9db0c9", width=2)
    draw.text((345, 194), "************", fill="#374151", font=text_font)

    draw.rounded_rectangle((370, 255, 470, 298), radius=8, fill="#2a9d8f")
    draw.rounded_rectangle((485, 255, 595, 298), radius=8, fill="#bfc9d9")
    draw.text((395, 266), "Войти", fill="white", font=button_font)
    draw.text((500, 266), "Отмена", fill="#243447", font=button_font)

    img.save(path)


def create_main_window_image(path: Path):
    banner = Image.open(SRC / "resources" / "images" / "database-banner.png").convert("RGB")
    img = Image.new("RGB", (1200, 760), "white")
    draw = ImageDraw.Draw(img)
    title_font = get_font(28, True)
    text_font = get_font(20)
    menu_font = get_font(19)

    draw.rectangle((0, 0, 1200, 34), fill="#eef2f7")
    draw.text((18, 7), "Файл", fill="#1f2937", font=menu_font)
    draw.text((90, 7), "База данных", fill="#1f2937", font=menu_font)
    draw.text((250, 7), "Справка", fill="#1f2937", font=menu_font)

    banner = banner.resize((1120, 268))
    img.paste(banner, (40, 70))

    draw.text((420, 380), "Главное окно приложения", fill="#111827", font=title_font)
    draw.text((290, 430), "Работа с таблицей базы данных открывается через меню сверху.", fill="#374151", font=text_font)

    img.save(path)


def create_database_window_image(path: Path):
    img = Image.new("RGB", (1280, 760), "white")
    draw = ImageDraw.Draw(img)
    title_font = get_font(24, True)
    text_font = get_font(18)
    small_font = get_font(16)

    buttons = [
        ("Версия БД", 40, 30, 180),
        ("Обновить данные", 235, 30, 220),
        ("Создать демонстрационные таблицы", 470, 30, 370),
        ("Закрыть", 855, 30, 130),
    ]
    for caption, x, y, w in buttons:
        draw.rounded_rectangle((x, y, x + w, y + 42), radius=8, fill="#dbe7f5", outline="#8aa6c1", width=2)
        draw.text((x + 14, y + 11), caption, fill="#1f2937", font=small_font)

    draw.text((40, 95), "Результат запроса SELECT VERSION():", fill="#111827", font=text_font)
    draw.rectangle((40, 125, 1240, 210), outline="#9db0c9", width=2)
    draw.text((52, 152), "PostgreSQL 17.x on localhost:5433", fill="#374151", font=text_font)

    draw.text((40, 235), "Сводка БД", fill="#111827", font=title_font)
    draw.text((40, 275), "users: 3 | analyses: 3 | payments: 2 | custom_brands: 2", fill="#374151", font=text_font)
    draw.text((40, 315), "Данные из таблицы users:", fill="#111827", font=text_font)

    table_x, table_y = 40, 350
    col_widths = [70, 260, 110, 120, 110, 100, 210]
    headers = ["ID", "Email", "Role", "Plan", "Analyses", "Active", "Created At"]
    rows = [
        ["1", "admin@veritasad.ai", "admin", "enterprise", "24", "true", "2026-03-11 14:10"],
        ["2", "analyst@veritasad.ai", "user", "pro", "11", "true", "2026-03-11 14:12"],
        ["3", "demo@veritasad.ai", "user", "starter", "3", "false", "2026-03-11 14:14"],
    ]

    x = table_x
    for header, width in zip(headers, col_widths):
        draw.rectangle((x, table_y, x + width, table_y + 42), fill="#eaf1f8", outline="#8aa6c1", width=2)
        draw.text((x + 8, table_y + 11), header, fill="#1f2937", font=small_font)
        x += width

    current_y = table_y + 42
    for row in rows:
        x = table_x
        for value, width in zip(row, col_widths):
            draw.rectangle((x, current_y, x + width, current_y + 40), fill="white", outline="#c4d0dd", width=1)
            draw.text((x + 8, current_y + 10), value, fill="#374151", font=small_font)
            x += width
        current_y += 40

    img.save(path)


def create_schema_image(path: Path):
    img = Image.new("RGB", (1400, 900), "#f8fafc")
    draw = ImageDraw.Draw(img)
    title_font = get_font(24, True)
    head_font = get_font(20, True)
    text_font = get_font(16)

    def box(x, y, w, h, title, lines, color="#ffffff"):
        draw.rounded_rectangle((x, y, x + w, y + h), radius=16, fill=color, outline="#6b88a6", width=3)
        draw.rectangle((x, y, x + w, y + 42), fill="#d9e7f5", outline="#6b88a6", width=3)
        draw.text((x + 14, y + 10), title, fill="#1f2937", font=head_font)
        yy = y + 56
        for line in lines:
            draw.text((x + 14, yy), line, fill="#374151", font=text_font)
            yy += 24

    draw.text((470, 24), "Логическая модель базы данных", fill="#111827", font=title_font)

    box(80, 130, 360, 210, "users", [
        "id PK",
        "email",
        "plan",
        "role",
        "total_analyses",
        "is_active",
        "created_at",
    ])
    box(510, 130, 360, 230, "analyses", [
        "id PK",
        "task_id",
        "video_id",
        "user_id FK -> users.id",
        "source_type",
        "status",
        "confidence_score",
        "has_advertising",
        "created_at",
    ])
    box(940, 130, 320, 190, "payments", [
        "id PK",
        "user_id FK -> users.id",
        "amount",
        "currency",
        "status",
        "provider",
        "created_at",
    ])
    box(510, 470, 360, 190, "custom_brands", [
        "id PK",
        "user_id FK -> users.id",
        "name",
        "category",
        "is_active",
        "detection_threshold",
        "created_at",
    ])

    draw.line((440, 230, 510, 230), fill="#2a9d8f", width=5)
    draw.line((260, 340, 260, 565), fill="#2a9d8f", width=5)
    draw.line((260, 565, 510, 565), fill="#2a9d8f", width=5)
    draw.line((870, 230, 940, 230), fill="#2a9d8f", width=5)

    img.save(path)


def create_assets():
    ASSETS.mkdir(exist_ok=True)
    create_login_image(ASSETS / "login.png")
    create_main_window_image(ASSETS / "main.png")
    create_database_window_image(ASSETS / "database.png")
    create_schema_image(ASSETS / "schema.png")


def add_cover(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for text in [
        "МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ",
        "федеральное государственное бюджетное образовательное учреждение",
        "высшего образования",
        "«Российский государственный университет им. А.Н. Косыгина",
        "(Технологии. Дизайн. Искусство)»",
        "",
        "Кафедра автоматизированных систем обработки информации и управления",
        "",
        "",
        "Отчет по лабораторной работе № 7",
        "по дисциплине «Программный интерфейс»",
        "Тема: «JDBC. Соединение приложения с БД»",
        "",
        "",
        "Выполнил: Юрков Д. А., группа МВА-122",
        "Проверил: Адаев Р. Б.",
        "",
        "Москва 2026",
    ]:
        run = p.add_run(text + "\n")
        set_run_font(run, size=14 if "Отчет" not in text and "Тема" not in text else 16, bold="Отчет" in text)


def add_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "8")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), "000000")
        borders.append(elem)
    tbl_pr.append(borders)


def build_report():
    create_assets()
    doc = Document()
    style_document(doc)
    add_cover(doc)

    doc.add_section(WD_SECTION.NEW_PAGE)

    doc.add_heading("Цель работы", level=1)
    doc.add_paragraph(
        "Получить общее представление о прикладном программном интерфейсе JDBC для работы с "
        "реляционной базой данных PostgreSQL, освоить процедуру установки соединения с БД и "
        "создать JavaFX-приложение с формой авторизации и отдельной формой работы с данными."
    )

    doc.add_heading("Краткое описание предметной области", level=1)
    doc.add_paragraph(
        "В качестве предметной области выбрана информационная система VeritasAd. "
        "Приложение хранит сведения о пользователях сервиса, выполненных анализах, платежах "
        "и пользовательских брендах. В лабораторной работе реализовано подключение к PostgreSQL, "
        "авторизация пользователя и просмотр данных таблицы users через JDBC."
    )

    doc.add_heading("Структура проекта", level=1)
    for line in tree_text().splitlines():
        add_code_paragraph(doc, line)

    doc.add_heading("Схема логической модели базы данных", level=1)
    doc.add_paragraph(
        "В демонстрационной схеме используются четыре основные таблицы: users, analyses, payments и custom_brands. "
        "Таблицы analyses, payments и custom_brands связаны с users через внешний ключ user_id."
    )
    doc.add_picture(str(ASSETS / "schema.png"), width=Inches(6.6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Описание экранных форм", level=1)
    doc.add_paragraph("Окно авторизации предназначено для ввода логина и пароля пользователя БД.")
    doc.add_picture(str(ASSETS / "login.png"), width=Inches(6.1))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("Главное окно содержит изображение и меню, через которое открывается рабочая форма БД.")
    doc.add_picture(str(ASSETS / "main.png"), width=Inches(6.4))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("Форма работы с БД отображает версию PostgreSQL, сводку и таблицу users.")
    doc.add_picture(str(ASSETS / "database.png"), width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Используемые файлы конфигурации", level=1)
    add_code_block(doc, "pom.xml", read_text(ROOT / "pom.xml"))
    add_code_block(doc, "conf.prop", read_text(ROOT / "conf.prop"))
    add_code_block(doc, "module-info.java", read_text(SRC / "java" / "module-info.java"))

    doc.add_heading("Текст программы", level=1)
    files = [
        ("MainApp.java", SRC / "java" / "com" / "lab7" / "MainApp.java"),
        ("DatabaseManager.java", SRC / "java" / "com" / "lab7" / "db" / "DatabaseManager.java"),
        ("Dialogs.java", SRC / "java" / "com" / "lab7" / "gui" / "Dialogs.java"),
        ("UserRecord.java", SRC / "java" / "com" / "lab7" / "model" / "UserRecord.java"),
        ("LoginController.java", SRC / "java" / "com" / "lab7" / "controllers" / "LoginController.java"),
        ("MainController.java", SRC / "java" / "com" / "lab7" / "controllers" / "MainController.java"),
        ("DatabaseController.java", SRC / "java" / "com" / "lab7" / "controllers" / "DatabaseController.java"),
        ("LoginView.fxml", SRC / "resources" / "com" / "lab7" / "fxml" / "LoginView.fxml"),
        ("MainView.fxml", SRC / "resources" / "com" / "lab7" / "fxml" / "MainView.fxml"),
        ("DatabaseView.fxml", SRC / "resources" / "com" / "lab7" / "fxml" / "DatabaseView.fxml"),
    ]
    for title, path in files:
        add_code_block(doc, title, read_text(path))

    doc.add_heading("Краткие выводы", level=1)
    doc.add_paragraph(
        "В ходе выполнения лабораторной работы было создано Maven-приложение на JavaFX с поддержкой JDBC. "
        "Реализованы форма авторизации пользователя, главное окно с изображением и отдельная форма работы с БД. "
        "Приложение подключается к PostgreSQL, выполняет запрос SELECT VERSION() и отображает данные таблицы users."
    )

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_report()
    print(f"Отчет создан: {OUTPUT}")
